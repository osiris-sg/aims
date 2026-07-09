#!/usr/bin/env python3
"""Build the AIMS User Manual (.docx) from annotated screenshots + step text.

Structure is data-driven (SECTIONS) so new sections/flows are easy to add.
Run: python3 build_manual.py
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

HERE = os.path.dirname(os.path.abspath(__file__))
ANN = os.path.join(HERE, "annotated")
NAVY = RGBColor(0x1A, 0x1A, 0x2E)
RED = RGBColor(0xE0, 0x31, 0x31)
GREY = RGBColor(0x66, 0x66, 0x66)

# ---- Manual content (extend this) -------------------------------------------
SECTIONS = [
    {
        "title": "Documents",
        "intro": "AIMS sales documents — Quotations, Sales Orders, Delivery Orders, "
                 "Invoices, and Credit/Debit Notes — are all created in the same editor "
                 "and follow the same steps; only the starting point and a few fields "
                 "differ. This section walks through creating each type, starting with a "
                 "Quotation.",
        "flows": [
            {
                "name": "The document toolbar",
                "image": "quot-toolbar.png",
                "intro": "Every document opens in the same editor. The toolbar across the top "
                         "provides the main actions (shown here on a Quotation):",
                "legend": [
                    "**Add** — start a new blank document of this type.",
                    "**Locate** — search for and open another existing document (by number, "
                    "customer name, P/O or D/O).",
                    "**Duplicate** — make a copy of the current document.",
                    "**Stock Card** — open the inventory item picker to add items (same as the "
                    "Add Item button on the line-items table).",
                    "**Print / PDF** — download or print the document.",
                    "**Preview** — view the finished document on your company letterhead.",
                    "**Confirm Quotation** — finalise the document and move it forward in the "
                    "pipeline. This label changes per document type (e.g. Confirm Order, Confirm "
                    "Invoice).",
                ],
                "note": "Also on the toolbar: the back arrow (←) returns to the list, "
                        "Previous / Next step through your documents, and the Saved indicator "
                        "shows changes are stored automatically — there is no separate Save button.",
            },
            {
                "name": "Creating a Quotation",
                "steps": [
                    ("Open the Quotation list and start a new quotation",
                     "In the left sidebar, expand **Sales** and click **Quotation**. At the "
                     "top-right of the Quotations List, click **Create Quotation**. A new draft "
                     "opens with an auto-generated quotation number.\n"
                     "Tip: the **Upload Quotation** button beside it lets AIMS read a PDF or image "
                     "and create the quotation for you automatically. This “extract” button is "
                     "named differently for each document type (for example, invoices use "
                     "**Import Invoices**), and a few document types don't offer it.",
                     "quot-01-list.png"),
                    ("Get to know the editor",
                     "The quotation editor has three areas: **①** the document details (number, "
                     "date, customer, terms), **②** the totals panel (gross, discount, GST, nett "
                     "total), and **③** the line-items table. Your work **auto-saves** — watch the "
                     "“Saved” indicator at the top-right.",
                     "quot-02-editor.png"),
                    ("Select the customer",
                     "Click the **Customer Code** field, type a few letters of the customer's name, "
                     "and press the search icon. Pick the customer from the **Locate Customer** "
                     "list — their address and contact fill in automatically.",
                     "quot-03-customer.png"),
                    ("Add line items",
                     "Click **Add Item** to open the Stock Card picker. Search by code, description "
                     "or category, then click an item to add it to the quotation. Use **Add Service** "
                     "instead for a non-stock charge (e.g. labour, delivery).",
                     "quot-04-additem.png"),
                    ("Enter quantity and price",
                     "In the line item, type the **Unit Cost** and **Qty**. The **Amount**, "
                     "**Sub-Total**, **GST** and **Nett Total** recalculate automatically. You can "
                     "apply a discount or change the tax setting in the totals panel on the right.",
                     "quot-05-filled.png"),
                    ("Preview, then print or confirm",
                     "Click **Preview** to see the finished quotation on your company letterhead. "
                     "From here use **Print / PDF** to download or print it, or **Confirm Quotation** "
                     "to finalise it and move it forward in the sales pipeline.",
                     "quot-06-preview.png"),
                ],
            },
            {
                "name": "Creating a Sales Order",
                "steps": [
                    ("Open Sales Order and start a new order",
                     "A **Sales Order** confirms a customer's order after they accept your quotation — "
                     "it's the internal record that drives delivery and invoicing. In the sidebar, open "
                     "**Sales › Sales Order** and click **Create Sales Order**. Fill it in **exactly "
                     "like a Quotation** (see 1.2) — select the customer and add line items. "
                     "**Upload Sales Order** creates one from a PDF/image.",
                     "so-01-list.png"),
                    ("Preview and finalise",
                     "Click **Preview** to see the finished Sales Order, then **Print / PDF** or "
                     "**Confirm Document** to finalise it.",
                     "so-02-preview.png"),
                ],
                "note": "The **Extract** button in the editor toolbar turns an accepted quotation "
                        "into this Sales Order — it opens “Extract Quotation to Sales Order”, lists "
                        "the customer's quotations, and copies the chosen quote's line items in so "
                        "you don't re-key them.",
            },
            {
                "name": "Creating a Delivery Order",
                "steps": [
                    ("Open Delivery Order and start a new DO",
                     "A **Delivery Order (DO)** is the packing slip that goes with goods delivered to "
                     "the customer — it records what was shipped. Open **Sales › Delivery Order** and "
                     "click **Create Delivery Order**. Fill it in like a Quotation; use the "
                     "**Delivery Address** tab to set where the goods are going. **Upload Delivery "
                     "Order** creates one from a file.",
                     "do-01-list.png"),
                    ("Preview and finalise",
                     "A Delivery Order is a packing slip — it lists **quantities, not prices** "
                     "(Item · Description · Quantity · UoM · Remarks). **Preview** shows the finished "
                     "DO; use **Print / PDF** or **Confirm Delivery**.",
                     "do-02-preview.png"),
                ],
                "note": "The **Extract** button in the editor toolbar turns a quotation into this "
                        "Delivery Order — “Extract Quotation to Delivery Order” lists the customer's "
                        "quotations; pick one to copy its items straight in.",
            },
            {
                "name": "Creating an Invoice",
                "steps": [
                    ("Open Invoice and start a new invoice",
                     "A **Tax Invoice** bills the customer for goods or services, including GST — it's "
                     "the document your customer pays against. Open **Sales › Invoice** and click "
                     "**Create Invoice**. Fill it in like a Quotation. Tip: inside the invoice editor, "
                     "**Extract from Quotation** pulls the customer and line items straight from an "
                     "existing quotation; **Import Invoices** creates invoices from uploaded files.",
                     "inv-01-list.png"),
                    ("Preview and issue",
                     "The finished **Tax Invoice** shows your GST registration number and totals. "
                     "Use **Print / PDF** to send it, or **Confirm Invoice** to issue it.",
                     "inv-02-preview.png"),
                ],
                "note": "The **Extract from Quotation** button in the editor toolbar does the same "
                        "for invoices — it copies the customer and line items from an existing "
                        "quotation straight into the invoice, so an accepted quote becomes an "
                        "invoice in one click.",
            },
            {
                "name": "Creating a Debit Note",
                "steps": [
                    ("Open Debit Note and start a new note",
                     "A **Debit Note** increases the amount a customer owes — for example to charge for "
                     "extra items or correct an under-billed invoice. Open **Sales › Debit Note** and "
                     "click **Create Debit Note**. Fill it in like a Quotation (it usually references "
                     "the original invoice).",
                     "dn-01-list.png"),
                    ("Preview and finalise",
                     "The finished **Debit Note**. Use **Print / PDF** or **Confirm Document**.",
                     "dn-02-preview.png"),
                ],
            },
            {
                "name": "Creating a Credit Note",
                "steps": [
                    ("Open Credit Note and start a new note",
                     "A **Credit Note** reduces the amount a customer owes — for returns, refunds, or "
                     "correcting an over-billed invoice. Open **Sales › Credit Note** and click "
                     "**Create Credit Note**. Fill it in like a Quotation (it usually references the "
                     "original invoice).",
                     "cn-01-list.png"),
                    ("Preview and finalise",
                     "The finished **Credit Note**. Use **Print / PDF** or **Confirm Document**.",
                     "cn-02-preview.png"),
                ],
            },
        ],
    },
    {
        "title": "Inventory",
        "intro": "Inventory has two layers: **Products** (the master record for an item — its "
                 "name, SKU key, category and default pricing) and **Inventory Items** (the actual "
                 "stock units of a product). This section covers adding both, plus the documents "
                 "that move stock — Purchases, Purchase Returns and Stock Adjustments.",
        "flows": [
            {
                "name": "Adding a Product",
                "steps": [
                    ("Open Products and start a new product",
                     "In the sidebar, open **Inventory › Products** and click **Add Product**. A "
                     "product is the master record for an item; every product needs a unique "
                     "**SKU key**. (To change the quantity of an existing product, use **Add Items** "
                     "on the Inventory Items page instead — see the next section.)",
                     "prod-01-list.png"),
                    ("Step 1 — Asset Creation",
                     "The form is a **3-step wizard** (Asset Creation → Additional Details → "
                     "Review). Fill the required fields: **Name**, **SKUKEY** (unique identifier), "
                     "**Category**, **Unit of Measure**, and the starting **Quantity**. Click "
                     "**Next**.",
                     "prod-02-form.png"),
                    ("Step 2 — Additional Details (optional)",
                     "Add optional details: a **Description**, **Cost Price** (what you pay — for "
                     "margin tracking), **Selling Price** (the default unit price when this product "
                     "appears on a document), custom price tiers, and a low-stock **Minimum "
                     "Quantity**. Click **Next**.",
                     "prod-03-details.png"),
                    ("Step 3 — Review and create",
                     "Review the details, optionally upload a product **Image**, then click "
                     "**Create Asset**.",
                     "prod-04-review.png"),
                    ("Done — the product is created",
                     "The new product appears at the top of the Products list with its SKU key, "
                     "category and stock. It's now selectable when adding items to quotations, "
                     "invoices and other documents.",
                     "prod-05-result.png"),
                ],
            },
            {
                "name": "Adding Inventory Items (stock)",
                "steps": [
                    ("Open Inventory Items and click Add Items",
                     "**Products** is the master record; **Inventory Items** are the individual "
                     "stock units of a product. To add stock to an existing product, open "
                     "**Inventory › Inventory Items** and click **Add Items**.",
                     "item-01-list.png"),
                    ("Fill in the stock details and save",
                     "Choose the **Asset** (the product), enter the **SKU** for the unit, the "
                     "**Quantity**, **Category**, **Location**, and **Status** (e.g. in stock), "
                     "then click **Save**. The stock is added to that product's on-hand quantity.",
                     "item-02-form.png"),
                ],
            },
            {
                "name": "Purchase Order (ordering stock)",
                "steps": [
                    ("Open Purchases and start a Purchase Order",
                     "Inventory documents move stock. A **Purchase Order (PO)** orders goods from a "
                     "supplier. Open **Inventory › Purchases** and click **Create Purchase Order** "
                     "(or **Upload Purchase Order** to create one from a PDF/image).",
                     "po-01-list.png"),
                    ("Select the supplier",
                     "Unlike sales documents, a PO is addressed to a **supplier**, not a customer. "
                     "Click the **Supplier Code** field's search icon to open the supplier picker.",
                     "po-02-editor.png"),
                    ("Pick the supplier",
                     "Choose the supplier from the list — their address and contact details fill in "
                     "automatically.",
                     "po-03-supplier.png"),
                    ("Add items and quantities",
                     "Click **Add Item** and pick products from the Stock Card. Each line shows the "
                     "**Cost Price** (what you pay the supplier); set the **Quantity** to order. The "
                     "totals update on the right.",
                     "po-04-filled.png"),
                    ("Preview, then receive the goods",
                     "**Preview** shows the finished PO on your letterhead. Use **Print / PDF** to "
                     "send it to the supplier, or **Receive** (top-right) to receive the goods into "
                     "stock — which increases your on-hand inventory.",
                     "po-05-preview.png"),
                ],
            },
            {
                "name": "Purchase Return (returning stock to a supplier)",
                "steps": [
                    ("Open Purchases Return and start a return",
                     "A **Purchase Return (PR)** returns goods to a supplier (e.g. faulty or wrong "
                     "items). Open **Inventory › Purchases Return** and click **Create Purchase "
                     "Return**.",
                     "pr-01-list.png"),
                    ("Select the supplier",
                     "Like a PO, a return is addressed to the **supplier**. Click the **Supplier "
                     "Code** search icon to open the picker.",
                     "pr-02-editor.png"),
                    ("Pick the supplier",
                     "Choose the supplier you're returning goods to.",
                     "pr-03-supplier.png"),
                    ("Add the returned items and preview",
                     "Add the items and quantities being returned, then **Preview**. Use **Print / "
                     "PDF** or **Confirm Document** to finalise.",
                     "pr-04-preview.png"),
                ],
            },
            {
                "name": "Stock Adjustment In (increase stock)",
                "steps": [
                    ("Open Stock Adjustment In and start one",
                     "A **Stock Adjustment In (SAI)** increases stock without a purchase — e.g. "
                     "found stock, a stock-count correction, or items returned back into inventory. "
                     "Open **Inventory › Stock Adjustment In** and click **Create Stock Adjustment "
                     "In**.",
                     "sai-01-list.png"),
                    ("Add the items being adjusted in",
                     "A stock adjustment has **no mandatory customer or supplier** — just the items. "
                     "Click **Add Item**, pick the product, and set the **Quantity** to add to "
                     "stock. (Supplier and W/O Number are optional.)",
                     "sai-02-editor.png"),
                    ("Preview and confirm",
                     "**Preview** the adjustment, then **Print / PDF** or **Confirm Stock Adjustment "
                     "In** to apply it — this **increases** the product's on-hand quantity.",
                     "sai-03-preview.png"),
                ],
            },
            {
                "name": "Stock Adjustment Out (decrease stock)",
                "steps": [
                    ("Open Stock Adjustment Out and start one",
                     "A **Stock Adjustment Out (SAO)** decreases stock without a sale — e.g. damage, "
                     "loss, samples, or write-offs. Open **Inventory › Stock Adjustment Out** and "
                     "click **Create Stock Adjustment Out**.",
                     "sao-01-list.png"),
                    ("Add the items being adjusted out",
                     "As with adjustment in, there's no customer or supplier — click **Add Item**, "
                     "pick the product, and set the **Quantity** to remove from stock.",
                     "sao-02-editor.png"),
                    ("Preview and confirm",
                     "**Preview**, then **Print / PDF** or **Confirm Stock Adjustment Out** to apply "
                     "it — this **reduces** the product's on-hand quantity.",
                     "sao-03-preview.png"),
                ],
            },
        ],
    },
    {
        "title": "Customers",
        "intro": "Customers are the companies and people you sell to. This section "
                 "covers adding a new customer to AIMS.",
        "flows": [
            {
                "name": "Creating a Customer",
                "steps": [
                    ("Open the Customers page and start a new customer",
                     "In the left sidebar, click **Customers**. At the top-right of the "
                     "Customers List, click the **Add Customer** button.",
                     "01-customers-list.png"),
                    ("Enter the customer's details",
                     "The Add Customer panel slides in from the right. Fill in the customer "
                     "information. **Name** is required; Email, Phone, Address and GST Reg. No. "
                     "are optional but recommended so they auto-fill on quotations and invoices.",
                     "02-add-customer-form.png"),
                    ("(Optional) Assign a salesman and contacts, then Save",
                     "Optionally pick a **Salesman** and add **Points of Contact** (the people "
                     "documents are addressed to, shown as “Attn To”). When the details are "
                     "correct, click **Save**.",
                     "03-filled-form.png"),
                    ("Done — the customer is created",
                     "The panel closes and the new customer appears at the top of the Customers "
                     "List, with an auto-generated Customer Code. You can now select this "
                     "customer when creating quotations, invoices and other documents.",
                     "04-result.png"),
                ],
            },
        ],
    },
    {
        "title": "Suppliers",
        "intro": "Suppliers are the vendors you buy from — they appear on Purchase Orders, "
                 "Purchase Returns and stock documents. Adding a supplier works just like adding a "
                 "customer.",
        "flows": [
            {
                "name": "Creating a Supplier",
                "steps": [
                    ("Open Suppliers and start a new supplier",
                     "In the left sidebar, click **Suppliers**, then click **Add Supplier** at the "
                     "top-right of the Suppliers List.",
                     "sup-01-list.png"),
                    ("Enter the supplier's details",
                     "Fill in the supplier information. **Name** is required; Email, Phone, Address "
                     "and GST Reg. No. are optional but recommended so they auto-fill on purchase "
                     "documents.",
                     "sup-02-form.png"),
                    ("Review and save",
                     "Check the details, then click **Save**.",
                     "sup-03-filled.png"),
                    ("Done — the supplier is created",
                     "The new supplier appears at the top of the Suppliers list with an "
                     "auto-generated Supplier Code. You can now select it when creating Purchase "
                     "Orders and other supplier documents.",
                     "sup-04-result.png"),
                ],
            },
        ],
    },
]

def set_cell_text(p, runs):
    """runs: list of (text, bold)."""
    for text, bold in runs:
        r = p.add_run(text)
        r.bold = bold

def md_bold_runs(text):
    """Split text on **bold** markers into (text, bold) runs."""
    parts = text.split("**")
    return [(p, i % 2 == 1) for i, p in enumerate(parts) if p]

def add_rich(p, text, color=None):
    """Add text to paragraph p, honouring **bold** and \\n line breaks."""
    for li, line in enumerate(text.split("\n")):
        if li > 0:
            p.add_run().add_break()
        for t, bold in md_bold_runs(line):
            r = p.add_run(t); r.bold = bold
            if color is not None:
                r.font.color.rgb = color

def add_image(doc, path, width=6.3):
    if os.path.exists(path):
        pic = doc.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(path, width=Inches(width))

def add_title_page(doc):
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("AIMS")
    r.font.size = Pt(54); r.font.bold = True; r.font.color.rgb = NAVY
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run("User Manual"); rs.font.size = Pt(24); rs.font.color.rgb = GREY
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rsub = sub.add_run("Asset & Inventory Management System"); rsub.font.size = Pt(13); rsub.font.color.rgb = GREY
    doc.add_page_break()

def style_base(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"; n.font.size = Pt(11)

def main():
    doc = Document()
    style_base(doc)
    add_title_page(doc)

    for si, section in enumerate(SECTIONS, 1):
        h = doc.add_heading(f"{si}. {section['title']}", level=1)
        h.runs[0].font.color.rgb = NAVY
        intro = doc.add_paragraph()
        for text, bold in md_bold_runs(section["intro"]):
            run = intro.add_run(text); run.bold = bold; run.font.color.rgb = GREY
        for fi, flow in enumerate(section["flows"], 1):
            fh = doc.add_heading(f"{si}.{fi}  {flow['name']}", level=2)
            fh.runs[0].font.color.rgb = NAVY

            if "steps" in flow:
                for n, (title, body, img) in enumerate(flow["steps"], 1):
                    sp = doc.add_paragraph()
                    badge = sp.add_run(f"Step {n}   "); badge.bold = True; badge.font.color.rgb = RED; badge.font.size = Pt(12)
                    th = sp.add_run(title); th.bold = True; th.font.size = Pt(12); th.font.color.rgb = NAVY
                    add_rich(doc.add_paragraph(), body)
                    add_image(doc, os.path.join(ANN, img))
                    doc.add_paragraph()
                if flow.get("note"):
                    add_rich(doc.add_paragraph(), "Note: " + flow["note"], color=GREY)
            else:  # reference flow: intro + image + numbered legend + note
                if flow.get("intro"):
                    add_rich(doc.add_paragraph(), flow["intro"])
                add_image(doc, os.path.join(ANN, flow["image"]))
                for i, item in enumerate(flow.get("legend", []), 1):
                    lp = doc.add_paragraph()
                    num = lp.add_run(f"{i}.  "); num.bold = True; num.font.color.rgb = RED
                    add_rich(lp, item)
                if flow.get("note"):
                    add_rich(doc.add_paragraph(), "Note: " + flow["note"], color=GREY)
                doc.add_paragraph()

    out = os.path.join(HERE, "AIMS-User-Manual.docx")
    doc.save(out)
    print(f"wrote {out}")

if __name__ == "__main__":
    main()
